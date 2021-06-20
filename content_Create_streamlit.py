import streamlit as st
import streamlit.components as stc
import pandas as pd
import numpy as np
import pathlib
import base64

#import plotly.express as px
#from plotly.subplots import make_subplots
#import plotly.graph_objects as go
#import matplotlib.pyplot as plt

import requests
import urllib
import pandas as pd
from requests_html import HTML
from requests_html import HTMLSession
import trafilatura
import altair as alt
from PIL import Image
from pathlib import Path
import time
import io
import os
import docx
from docx import Document
from fake_useragent import UserAgent
from bs4 import BeautifulSoup
import re
import sys

from docx.shared import Inches, Cm
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_UNDERLINE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
#import nltk
#nltk.download('punkt')
#from nltk.tokenize import sent_tokenize

def img_to_bytes(img_path):
    img_bytes = Path(img_path).read_bytes()
    encoded = base64.b64encode(img_bytes).decode()
    return encoded

header_html = "<img src='data:image/png;base64,{}' class='img-fluid'>".format(
    img_to_bytes("DSLogo.png")
)
st.markdown(
    header_html, unsafe_allow_html=True,
)
#image = Image.open('C:\\Users\\Darcey\\Downloads\\DeepSphere Logo.jpg')

#st.image(image)

#LOGO_IMAGE = "C:\\Users\\Darcey\\Downloads\\DeepSphere Logo.jpg"
#
#st.markdown(
#    """
#    <style>
#    .container {
#        display: flex;
#    }
#    .logo-text {
#        font-weight:70 !important;
#        font-size:10px !important;
#        color: #f9a01b !important;
#        padding-top: 75px !important;
#    }
#    .logo-img {
#        float:right;
#    }
#    </style>
#    """,
#    unsafe_allow_html=True
#)
#
#st.markdown(
#    f"""
#    <div class="container">
#        <img class="logo-img" src="data:image/png;base64,{base64.b64encode(open(LOGO_IMAGE, "rb").read()).decode()}">
#        <p class="logo-text">Logo Much ?</p>
#    </div>
#    """,
#    unsafe_allow_html=True
#)


st.title("Content Creation for the Given Topic using **_Web Scraping_** and **_NLP_**")
st.sidebar.title("Content Creation for the Given Topic using Web Scraping and NLP")
#st.markdown("This application is to extract URLs and text content related to the given topic:")
st.markdown("***")
st.sidebar.markdown("This application is to extract URLs and text content for the given topic")

def get_source(url):
    """Return the source code for the provided URL. 

    Args: 
        url (string): URL of the page to scrape.

    Returns:
        response (object): HTTP response object from requests_html. 
    """

    try:
        session = HTMLSession()
        response = session.get(url)
        return response

    except requests.exceptions.RequestException as e:
        print(e)
        
def scrape_google(query):

    query = urllib.parse.quote_plus(query)
    response = get_source("https://www.google.co.in/search?q=" + query)

    links = list(response.html.absolute_links)
    google_domains = ('https://www.google.', 
                      'https://google.', 
                      'https://webcache.googleusercontent.', 
                      'http://webcache.googleusercontent.', 
                      'https://policies.google.',
                      'https://support.google.',
                      'https://maps.google.',
                      'https://www.youtube.')

    for url in links[:]:
        if url.startswith(google_domains):
            links.remove(url)

    return links

def text_downloader(raw_text):
	b64 = base64.b64encode(raw_text.encode()).decode()
	new_filename = "new_text_file_{}_.docx".format(timestr)
	st.markdown("#### Download File ###")
	href = f'<a href="data:file/docx;base64,{b64}" download="{new_filename}">Click Here!!</a>'
	st.markdown(href,unsafe_allow_html=True)
    
class FileDownloader(object):
	"""docstring for FileDownloader
	>>> download = FileDownloader(data,filename,file_ext).download()
	"""
	def __init__(self, data,filename='myfile',file_ext='docx'):
		super(FileDownloader, self).__init__()
		self.data = data
		self.filename = filename
		self.file_ext = file_ext

	def download(self):
		b64 = base64.b64encode(self.data.encode()).decode()
		new_filename = "{}_{}_.{}".format(self.filename,timestr,self.file_ext)
		st.markdown("#### Download File ###")
		href = f'<a href="data:file/{self.file_ext};base64,{b64}" download="{new_filename}">Click Here!!</a>'
		st.markdown(href,unsafe_allow_html=True)

def download_link(object_to_download, download_filename, download_link_text):
    """
    Generates a link to download the given object_to_download.
    object_to_download (str, pd.DataFrame):  The object to be downloaded.
    download_filename (str): filename and extension of file. e.g. mydata.csv, some_txt_output.txt
    download_link_text (str): Text to display for download link.
    Examples:
    download_link(YOUR_DF, 'YOUR_DF.csv', 'Click here to download data!','required_question_type')
    download_link(YOUR_STRING, 'YOUR_STRING.txt', 'Click here to download your text!','required_question_type')
    """
#     if isinstance(object_to_download,pd.DataFrame):
#         object_to_download = object_to_download.to_csv(index=False)
    doc = Document()
#    with open("C:\\Users\\Darcey\\Downloads\\model_output.txt", 'r', encoding='utf-8') as file:
    doc.add_paragraph(object_to_download)
    doc_to_save = doc.save(str(Topic)+".docx")
#    object_to_download1 = ""

    return doc_to_save

#Topic = st.text_input('Input the topic here:')

def scrape_google_all(Topic):

    Topic = urllib.parse.quote_plus(Topic) # Format into URL encoding
    number_result = 10

    ua = UserAgent()

    google_url = "https://www.google.com/search?q=" + Topic + "&num=" + str(number_result)
    response = requests.get(google_url, {"User-Agent": ua.random})
    soup = BeautifulSoup(response.text, "html.parser")

    result_div = soup.find_all('div', attrs = {'class': 'ZINbbc'})

    links = []
    titles = []
    descriptions = []
    for r in result_div:
        # Checks if each element is present, else, raise exception
        try:
            link = r.find('a', href = True)
            title = r.find('div', attrs={'class':'vvjwJb'}).get_text()
            description = r.find('div', attrs={'class':'s3v9rd'}).get_text()

            # Check to make sure everything is present before appending
            if link != '' and title != '' and description != '': 
                links.append(link['href'])
                titles.append(title)
                descriptions.append(description)
        # Next loop if one element is not present
        except:
            continue
    return links

def Extract_Ranked_urls(links):
    to_remove = []
    clean_links = []
    for i, l in enumerate(links):
        clean = re.search('\/url\?q\=(.*)\&sa',l)

        # Anything that doesn't fit the above pattern will be removed
        if clean is None:
            to_remove.append(i)
            continue
        clean_links.append(clean.group(1))

    # Remove the corresponding titles & descriptions
#    for x in to_remove:
#        del titles[x]
#        del descriptions[x]
    return clean_links

def Extract_urls(Topic):
    df = pd.DataFrame(scrape_google(Topic), columns = ['link'])
    return df

def Extract_Contents(clean_links):
    list2 = []
    i=1
    for url in clean_links:
        downloaded = trafilatura.fetch_url(url)
        trafilatura.extract(downloaded)
        # outputs main content and comments as plain text ...
        list1 = trafilatura.extract(downloaded, include_comments=False)
        # outputs main content without comments as XML ...
        list2.append("\n")
        list2.append("---------------------------------------------------------------------------------------------------------------------")
        list2.append("\n")
        list2.append("**Content Set #")
        list2.append(str(i))
        list2.append("**")
        list2.append("\n")
        list2.append("\n")
        list2.append("URL #")
        list2.append(str(i))
        list2.append(":    ")
        list2.append(url)
        list2.append("\n")
        list2.append("---------------------------------------------------------------------------------------------------------------------")
        list2.append("\n")
        list2.append("\n")
        list2.append(list1)
        list2.append("\n")
        list2.append("---------------------------------------------------------------------------------------------------------------------")
        list2.append("\n")
        list2.append("---------------------------------------------------------------------------------------------------------------------")
        list2.append("\n")
        list2.append("\n")
        i+=1
        list3 = ''.join(filter(None, list2))
    return list3

def View_Extracted_Contents(list3):
    Extracted_Contents = list3
#    data = [content.strip() for content in Extracted_Contents.splitlines() if content]
#    data1 = ''.join([str(elem) for elem in data])
    return Extracted_Contents

def para_correct(list3):
    data = [content.strip() for content in list3.splitlines() if content]
    data1 = ''.join([str(elem) for elem in data])
    return data1

##################### To add page number in the footer ###############################################

def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' of '
    of_run._r.append(t2)

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')

    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')

    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)
    
##################### To add page number in the footer ###############################################


def main():
    Topic = st.text_input('Input the topic here and press ENTER:')
    i=1
#if len(Topic)>0:
    if st.sidebar.button("Extract URLs for the given topic"):
        with st.spinner("Extracting..."):
            links = scrape_google_all(Topic)
            clean_links = Extract_Ranked_urls(links)
            st.write("Below are the top URLs to extract content:")
            for x in clean_links:
                st.write(x)
    st.sidebar.markdown("*******************************")
    if st.sidebar.button("Download Contents from URLs"):
#    if text is not None:
        with st.spinner("Downloading..."):
#            i=1
            links = scrape_google_all(Topic)
            clean_links = Extract_Ranked_urls(links)
            list3 = Extract_Contents(clean_links)
#            data1 = para_correct(list3)
            data = [content.strip() for content in list3.splitlines() if content]
            data1 = '\\n\n'.join(f"{row}\n" for row in data)
        
############ Converting listof urls to dataframe and then to tuple to create the table in word document##########

            df = pd.DataFrame(clean_links, columns = ['urls'])
            df['url_rank'] = np.arange(len(df)) + 1
            df1 = df[['url_rank', 'urls']]
#            print (df1)
            datat = tuple(df1.itertuples(index=False, name=None))
    

#            print(datat)
            doc = docx.Document()

##################### To add logo ###############################################
            ##### add logo in normal header
#            doc = docx.Document()
            #p = doc.add_paragraph()
            #r = p.add_run()
            #r.add_picture('C:\\Users\\Darcey\\Downloads\\DeepSphere Logo.jpg', width=Inches(1.5))

            
#            sections2 = doc.sections
#            for section in sections2:
#                section.top_margin = Cm(0.5)
#                section.bottom_margin = Cm(0.5)
#                section.left_margin = Cm(1)
#                section.right_margin = Cm(1)
            ##### add logo in Zoned header
            
            logo_path = 'DSLogo.png'    # Path of the image file
            section = doc.sections[0]   # Create a section
            sec_header = section.header   # Create header 
            header_tp = sec_header.add_paragraph()  # Add a paragraph in the header, you can add any anything in the paragraph
            header_run = header_tp.add_run()   # Add a run in the paragraph. In the run you can set the values 
            header_run.add_picture(logo_path, width=Inches(1.3))  # Add a picture and set width.
            #rml_header = "\t Applied Artificial Intelligence for Schools Content \t Generation by Topic \t"
            header_run.add_text("\n                                                                                                 ")
            header_run.add_text("Applied Artificial Intelligence for Schools Content Generation by Topic")
            header_run.add_text("\n__________________________________________________________________________________")
            header_run.font.size =  Pt(14)
    
            #section1 = doc.sections[0]
            #header = section1.header
            #htable=header.add_table(1, 2, Inches(6))
            #htab_cells=htable.rows[0].cells
            #ht0=htab_cells[0].add_paragraph()
            #kh=ht0.add_run()
            #kh.add_picture('C:\\Users\\Darcey\\Downloads\\DeepSphere Logo.jpg', width=Inches(1.2))
            #ht1=htab_cells[1].add_paragraph('Applied Artificial Intelligence for Schools Content Generation by Topic')
            #ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
#            ht1.style.font.size = Pt(13)
#            ht1.style.font.bold = True
            #paragraph = header.paragraphs[0]
            #run1 = paragraph.add_run()
            #run1.add_picture("C:\\Users\\Darcey\\Downloads\\DeepSphere Logo.jpg", width=Inches(1))
            #run1.text = 'Applied Artificial Intelligence for Schools Content Generation by Topic'
            #run1.alignment=WD_ALIGN_PARAGRAPH.LEFT
            #run1.font.color.rgb = RGBColor(0, 0, 0)
            #run1.font.size = Pt(13)
            #run1.font.bold = True
#            run1.font.underline = True
            doc.add_paragraph('')
#            doc.add_paragraph('')

##################### To add footer with page number ###############################################

            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "© DeepSphere.AI | Confidential and Proprietary |Not for Distribution \t"
            add_page_number(doc.sections[0].footer.paragraphs[0])
#            heading = doc.add_heading('\t Applied Artificial Intelligence for Schools Content Generation by Topic \t', 0)
#            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
#            heading.style.font.color.rgb = RGBColor(0, 0, 0)
#            heading.style.font.size = Pt(18)
#            heading.style.font.bold = True
            #heading.style.font.underline = True

##################### To add given topic as sub header ###############################################

            topic1 = "Topic: " + Topic
            #topic = doc.add_heading("Topic: ", 1)
            topic = doc.add_heading(topic1, 1)
            doc.add_paragraph('')
            topic.style.font.color.rgb = RGBColor(0, 0, 0)
            topic.style.font.size = Pt(14)
            topic.style.font.bold = True
            
##################### To add table with url and its ranking ###############################################

            table = doc.add_table(rows=1, cols=2)
            row = table.rows[0].cells
            row[0].text = 'URL Rank'
            row[1].text = 'URLs'
            for url_rank, urls in datat:
                row = table.add_row().cells
                row[0].text = str(url_rank) + "                               "
                row[1].text = urls

            table.style = 'Colorful List'
            table.autofit = False
            table.allow_autofit = False
            cell = table.rows[0].cells[0]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell1 = table.rows[0].cells[1]
            cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

##################### To export/save the document ###############################################


#            for x in list2:
    
#    doc = docx.Document()
            doc.add_page_break()
            sent2 = sent_tokenize(list3)
#            doc = docx.Document()
            for ss in sent2:
                doc.add_paragraph(ss)
            doc.save("Model Output - " + str(Topic)+".docx")
#            doc.add_paragraph(data1)
#    para.paragraph_format.space_before = Inches(0.5)
#    para.paragraph_format.space_after = Inches(0.5)

#doc.add_heading('Applied Artificial Intelligence for Schools Content Generation by Topic', 0)
#            doc.save(str(Topic)+".docx")


#            doc = Document()
#            doc.add_paragraph(clean_links)
#            doc.add_paragraph(data1)
#            docx = Document(io.BytesIO(requests.get(doc).content))
#            b64 = base64.b64encode(docx)  # some strings <-> bytes conversions necessary here
#            href = f'<a href="data:file/docx;base64,{b64}">Download docx file</a>'
#            st.markdown(href, unsafe_allow_html=True)
#            doc.paragraph_format.space_after = Inches(1.0)
#            try:
#                doc.save(str(Topic)+".docx")
#            except:
#                print(sys.exc_info()[0], "occurred.")
        st.markdown("Download Complete")
    st.sidebar.markdown("*******************************")
    if st.sidebar.checkbox("View the Extracted Contents"):
        with st.spinner("Downloading the Contents..."):
            links = scrape_google_all(Topic)
            clean_links = Extract_Ranked_urls(links)
            list3 = Extract_Contents(clean_links)
            Extracted_Contents = View_Extracted_Contents(list3)
            data = [content.strip() for content in Extracted_Contents.splitlines() if content]
            for x in data:
                st.write(x)
            list2 = []
            for url in clean_links:
                downloaded = trafilatura.fetch_url(url)
                trafilatura.extract(downloaded)
                # outputs main content and comments as plain text ...
                list1 = trafilatura.extract(downloaded, include_comments=False)
                st.write("***************************************")
                st.write(url)
                if list1 is None:
                    st.write("Contents not available")
                else:
                    st.write("Contents available")
                ua = UserAgent()
                response = requests.get(url, {"User-Agent": ua.random})

                st.write("Response Code: ", response.status_code)
#        if not st.sidebar.checkbox("View the Extracted Contents"):
#            with st.spinner("Fetching the link to download..."):
#            df = Extract_urls(Topic)
#            list3 = Extract_Contents(df)
#            st.markdown(download_link(list3, 'model_output.txt', 'Click here to download the extracted text'),unsafe_allow_html=True)
#        View_Option = st.sidebar.radio("To view or Download Content: ",
#                     ('View', 'Download'))
#        if (View_Option == 'View'):
#            st.write(list3)
#        elif (View_Option == 'Download'):
#            st.markdown(download_link(list3, 'model_output.txt', 'Click here to download the extracted text'),unsafe_allow_html=True)
main()
